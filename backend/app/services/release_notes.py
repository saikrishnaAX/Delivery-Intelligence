"""Generate release notes from tickets moved to the Released section in Asana."""

from __future__ import annotations

import io
import re
from datetime import datetime, timedelta, time
from typing import Literal

from dateutil import parser as date_parser
from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session, joinedload

from app.config import get_settings
from app.integrations.asana import AsanaClient
from app.models import AsanaProject, Ticket, TicketSectionMove
from app.services.section_tracking import backfill_project_releases
from app.services.section_utils import is_released_section
from app.services.ticket_parser import (
    clean_title_for_release,
    infer_module_affected,
    redact_workshop_names,
)

from app.services.release_notes_polish import polish_release_items

settings = get_settings()

ReleaseCategory = Literal["enhancement", "performance", "bug", "security"]

CATEGORY_LABELS = {
    "enhancement": "Enhancements / Requirements Implemented",
    "performance": "Performance Improvements",
    "bug": "Bugs Fixed",
    "security": "Security Enhancements",
}

CATEGORY_EMOJI = {
    "enhancement": "⭐",
    "performance": "🚀",
    "bug": "🐞",
    "security": "🔒",
}

TITLE_EMOJIS = [
    (("registration", "vehicle"), "🚗"),
    (("parts", "inward"), "🔧"),
    (("refund",), "💸"),
    (("invoice",), "🧾"),
    (("sms", "whatsapp"), "📊"),
    (("quick pay", "payment"), "💳"),
    (("estimation", "estimate"), "💰"),
    (("notification", "cep", "gms"), "🔔"),
    (("bilingual", "arabic"), "🌐"),
    (("mrp", "price"), "💰"),
    (("service type",), "🛠️"),
    (("kuwait",), "🧾"),
    (("sales register", "export"), "📊"),
    (("duplicate",), "📊"),
    (("vendor", "payment"), "💳"),
    (("tax", "gst", "igst", "sez"), "🧾"),
    (("barcode",), "🖨️"),
    (("search", "labour"), "🔍"),
    (("performance", "slow", "optimization"), "🚀"),
    (("bug", "fix", "issue", "resolved"), "🐞"),
]


def _parse_dt(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return date_parser.parse(value).replace(tzinfo=None)
    except (ValueError, TypeError):
        return None


def _window_from_lookback(lookback_days: int, date_to: datetime | None = None) -> tuple[datetime, datetime]:
    end_day = (date_to or datetime.utcnow()).date()
    start_day = end_day - timedelta(days=max(lookback_days - 1, 0))
    return datetime.combine(start_day, time.min), datetime.combine(end_day, time.max)


def _is_released_section(name: str | None) -> bool:
    return is_released_section(name)


def _pick_emoji(title: str, category: ReleaseCategory) -> str:
    lower = title.lower()
    for keywords, emoji in TITLE_EMOJIS:
        if all(k in lower for k in keywords):
            return emoji
    return CATEGORY_EMOJI[category]


def _classify_ticket(title: str, description: str, type_raw: str | None) -> ReleaseCategory:
    blob = f"{title} {description} {(type_raw or '')}".lower()
    if any(w in blob for w in ("security", "vulnerability", "unauthorized", "access control", "otp", "link security")):
        return "security"
    if type_raw and "bug" in type_raw.lower():
        return "bug"
    perf_words = ("performance", "optimization", "optimiz", "slow", "large data", "export via email", "load time")
    if any(w in blob for w in perf_words):
        return "performance"
    bug_words = ("fix", "fixed", "resolved", "issue", "bug", "incorrect", "duplicate", "not working")
    if any(w in blob for w in bug_words) and "enhancement" not in blob:
        return "bug"
    return "enhancement"


def _extract_whats_new(description: str) -> list[str]:
    lines = [ln.strip() for ln in description.splitlines() if ln.strip()]
    bullets: list[str] = []
    for line in lines:
        lower = line.lower()
        if lower.startswith(("issue ", "expected behavior", "actual behavior", "steps to reproduce")):
            continue
        cleaned = re.sub(r"^[\-\*•\d\.\)]+\s*", "", line).strip()
        if not cleaned or len(cleaned) < 8:
            continue
        if cleaned.lower().startswith(("garage name", "workshop name", "ax id")):
            continue
        if lower.startswith(("summary", "what's new", "impact", "note:")):
            cleaned = cleaned.split(":", 1)[-1].strip()
            if len(cleaned) >= 8:
                bullets.append(cleaned)
            continue
        if line.startswith(("-", "*", "•")) or re.match(r"^\d+[\.\)]", line):
            bullets.append(cleaned)
        elif any(lower.startswith(p) for p in ("added ", "introduced ", "enabled ", "updated ", "removed ", "fixed ", "supports ", "allows ", "improved ")):
            bullets.append(cleaned)
    if not bullets:
        for line in lines[:6]:
            if len(line) > 20 and not line.lower().startswith(("issue ", "expected")):
                bullets.append(line)
                if len(bullets) >= 4:
                    break
    return bullets[:8]


def _build_summary(description: str, title: str) -> str:
    text = description.strip()
    if not text:
        return f"Changes delivered for {title}."
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    for para in paragraphs:
        lower = para.lower()
        if lower.startswith(("issue ", "expected behavior", "actual behavior")):
            continue
        sentence = para.split("\n")[0].strip()
        if len(sentence) > 30:
            return sentence[:400]
    first = text.split("\n")[0].strip()
    return first[:400] if len(first) > 20 else f"Updates related to {title}."


def _build_impact(category: ReleaseCategory, summary: str) -> str:
    if category == "bug":
        return "Resolves reported issues and improves system reliability for affected workflows."
    if category == "performance":
        return "Improves system responsiveness and reduces load when handling large datasets."
    if "setting" in summary.lower() or "configur" in summary.lower():
        return "Provides configurable control while keeping default behavior unchanged for other users."
    return "Improves usability and operational efficiency for the affected module workflows."


def _format_release_category(category: ReleaseCategory) -> str:
    if category == "bug":
        return "Bug Fix"
    if category == "performance":
        return "Performance Improvement"
    if category == "security":
        return "Security Enhancement"
    return "Enhancement"


def _build_executive_summary(sections: dict, window_end: datetime) -> dict:
    counts = {key: len(sections.get(key) or []) for key in ("enhancement", "security", "performance", "bug")}
    total = sum(counts.values())
    parts: list[str] = []
    if counts["enhancement"]:
        parts.append(f"{counts['enhancement']} product improvement{'s' if counts['enhancement'] != 1 else ''}")
    if counts["bug"]:
        parts.append(f"{counts['bug']} customer issue{'s' if counts['bug'] != 1 else ''} resolved")
    if counts["security"]:
        parts.append(f"{counts['security']} security hardening update{'s' if counts['security'] != 1 else ''}")
    if counts["performance"]:
        parts.append(f"{counts['performance']} performance gain{'s' if counts['performance'] != 1 else ''}")

    headline = f"{total} release item{'s' if total != 1 else ''} shipped" if total else "No releases in this window"
    subheadline = ", ".join(parts) if parts else "Nothing moved to Released during the selected period."

    highlights: list[dict] = []
    for key in ("enhancement", "security", "performance", "bug"):
        for item in (sections.get(key) or [])[:2]:
            highlights.append({
                "title": item.get("title"),
                "category": key,
                "benefit": item.get("impact_benefit") or item.get("impact") or item.get("summary"),
            })
        if len(highlights) >= 5:
            break

    return {
        "headline": headline,
        "subheadline": subheadline,
        "counts": counts,
        "total": total,
        "highlights": highlights[:5],
        "release_month": window_end.strftime("%B %Y"),
    }


class ReleaseNotesService:
    def __init__(self, db: Session, project_gid: str | None = None):
        self.db = db
        self.project_gid = project_gid
        self.asana = AsanaClient()

    def _project(self) -> AsanaProject | None:
        if not self.project_gid:
            return None
        return self.db.query(AsanaProject).filter(AsanaProject.gid == self.project_gid).first()

    def _collect_moved_tickets(
        self,
        window_start: datetime,
        window_end: datetime,
        sprint_sheet_id: int | None = None,
    ) -> list[dict]:
        project = self._project()
        if not project:
            return []

        seen: set[int] = set()
        items: list[dict] = []

        moves = (
            self.db.query(TicketSectionMove)
            .options(joinedload(TicketSectionMove.ticket))
            .join(Ticket)
            .filter(
                Ticket.project_id == project.id,
                TicketSectionMove.moved_at >= window_start,
                TicketSectionMove.moved_at <= window_end,
            )
            .order_by(TicketSectionMove.moved_at.desc())
            .all()
        )
        for move in moves:
            if not _is_released_section(move.to_section):
                continue
            ticket = move.ticket or self.db.query(Ticket).filter(Ticket.id == move.ticket_id).first()
            if not ticket or ticket.id in seen:
                continue
            seen.add(ticket.id)
            items.append(self._ticket_to_item(ticket, move.moved_at, move))

        released_tickets = (
            self.db.query(Ticket)
            .options(joinedload(Ticket.module))
            .filter(
                Ticket.project_id == project.id,
                Ticket.released_at.isnot(None),
                Ticket.released_at >= window_start,
                Ticket.released_at <= window_end,
            )
            .all()
        )
        for ticket in released_tickets:
            if ticket.id in seen:
                continue
            if ticket.module and not _is_released_section(ticket.module.name):
                continue
            seen.add(ticket.id)
            items.append(self._ticket_to_item(ticket, ticket.released_at))

        if sprint_sheet_id:
            from app.models import SprintSheetRow
            sheet_rows = (
                self.db.query(SprintSheetRow)
                .filter(
                    SprintSheetRow.sheet_id == sprint_sheet_id,
                    SprintSheetRow.sheet_status == "released",
                )
                .all()
            )
            for srow in sheet_rows:
                ticket = srow.ticket or self.db.query(Ticket).filter(Ticket.id == srow.ticket_id).first()
                if not ticket or ticket.id in seen:
                    continue
                moved = ticket.released_at or srow.updated_at or datetime.utcnow()
                if moved < window_start or moved > window_end:
                    continue
                seen.add(ticket.id)
                item = self._ticket_to_item(ticket, moved, sheet_row=srow)
                items.append(item)

        items.sort(key=lambda x: x["moved_at"], reverse=True)
        return items

    def _ticket_to_item(
        self,
        ticket: Ticket,
        moved_at: datetime,
        move: TicketSectionMove | None = None,
        sheet_row=None,
    ) -> dict:
        workshop = ticket.workshop_name
        title = clean_title_for_release(ticket.title, workshop)
        description = redact_workshop_names(ticket.description or "", workshop)

        if sheet_row and sheet_row.row_data:
            rd = sheet_row.row_data
            if rd.get("release"):
                description = f"{description}\n\nRelease notes: {rd['release']}".strip()
            if rd.get("team"):
                pass  # team stored separately if needed

        category = _classify_ticket(title, description, ticket.asana_type_raw)
        whats_new = _extract_whats_new(description)
        if sheet_row and sheet_row.row_data and sheet_row.row_data.get("release"):
            release_note = str(sheet_row.row_data["release"]).strip()
            if release_note and release_note not in whats_new:
                whats_new = [release_note, *whats_new][:8]
        summary = redact_workshop_names(_build_summary(description, title), workshop)
        module_affected = infer_module_affected(title, description)
        return {
            "ticket_id": ticket.id,
            "asana_gid": ticket.asana_gid,
            "title": title,
            "category": category,
            "release_category": _format_release_category(category),
            "module_affected": module_affected,
            "summary": summary,
            "whats_new": whats_new,
            "impact": _build_impact(category, summary),
            "emoji": _pick_emoji(title, category),
            "moved_at": moved_at.isoformat(),
            "assignee": ticket.assignee,
            "asana_url": ticket.asana_url,
            "from_section": move.from_section if move else None,
        }

    async def build(
        self,
        lookback_days: int = 2,
        date_from: str | None = None,
        date_to: str | None = None,
        sprint_name: str | None = None,
        refresh_releases: bool = False,
    ) -> dict:
        project = self._project()
        if date_from and date_to:
            parsed_start = _parse_dt(date_from)
            parsed_end = _parse_dt(date_to)
            if parsed_start and parsed_end:
                window_start = datetime.combine(parsed_start.date(), time.min)
                window_end = datetime.combine(parsed_end.date(), time.max)
                if window_end < window_start:
                    window_start, window_end = window_end, window_start
                    window_start = datetime.combine(window_start.date(), time.min)
                    window_end = datetime.combine(window_end.date(), time.max)
            else:
                window_start, window_end = _window_from_lookback(lookback_days)
        else:
            window_start, window_end = _window_from_lookback(lookback_days)

        sprint_sheet_id = None
        if sprint_name and project:
            from app.models import SprintSheet
            sheet = (
                self.db.query(SprintSheet)
                .filter(SprintSheet.project_id == project.id, SprintSheet.name == sprint_name)
                .first()
            )
            if sheet:
                sprint_sheet_id = sheet.id

        if refresh_releases and project and self.asana.is_configured:
            await backfill_project_releases(
                self.db, project.id, self.asana, only_missing=True
            )
            self.db.commit()

        items = self._collect_moved_tickets(window_start, window_end, sprint_sheet_id)
        items = polish_release_items(items)
        grouped = {
            "enhancement": [i for i in items if i["category"] == "enhancement"],
            "performance": [i for i in items if i["category"] == "performance"],
            "bug": [i for i in items if i["category"] == "bug"],
            "security": [i for i in items if i["category"] == "security"],
        }
        executive_summary = _build_executive_summary(grouped, window_end)
        return {
            "project_name": project.name if project else None,
            "project_gid": self.project_gid,
            "released_section": settings.asana_released_section_name,
            "window_start": window_start.isoformat(),
            "window_end": window_end.isoformat(),
            "lookback_days": lookback_days,
            "release_date": window_end.strftime("%d %B %Y"),
            "document_title": f"Product Release — {window_end.strftime('%d %B %Y')}",
            "total_items": len(items),
            "sections": grouped,
            "items": items,
            "executive_summary": executive_summary,
            "asana_live": self.asana.is_configured,
            "source": "database",
        }

    def build_docx(self, payload: dict) -> bytes:
        doc = Document()
        title = doc.add_paragraph()
        run = title.add_run(payload.get("document_title") or "Product Release")
        run.bold = True
        run.font.size = Pt(20)

        date_line = doc.add_paragraph()
        date_run = date_line.add_run(f"Release date: {payload['release_date']}")
        date_run.font.size = Pt(11)

        summary = payload.get("executive_summary") or {}
        if summary.get("headline"):
            doc.add_paragraph("")
            exec_p = doc.add_paragraph()
            er = exec_p.add_run(summary["headline"])
            er.bold = True
            er.font.size = Pt(12)
            if summary.get("subheadline"):
                doc.add_paragraph(summary["subheadline"])

        doc.add_paragraph("")

        section_order = ("enhancement", "security", "performance", "bug")
        for category in section_order:
            section_items = payload["sections"].get(category) or []
            if not section_items:
                continue
            heading = doc.add_paragraph()
            hrun = heading.add_run(CATEGORY_LABELS[category])
            hrun.bold = True
            hrun.font.size = Pt(14)
            doc.add_paragraph("")

            for item in section_items:
                item_title = doc.add_paragraph()
                tr = item_title.add_run(item["title"])
                tr.bold = True
                tr.font.size = Pt(11)

                if category == "bug":
                    fix_text = item.get("fix") or item.get("summary") or ""
                    if fix_text:
                        fix = doc.add_paragraph()
                        fr = fix.add_run("Resolution: ")
                        fr.bold = True
                        fix.add_run(fix_text)
                    doc.add_paragraph("")
                    continue

                benefit = item.get("impact_benefit") or item.get("impact") or ""
                if benefit:
                    impact = doc.add_paragraph()
                    ir = impact.add_run("Business impact: ")
                    ir.bold = True
                    impact.add_run(benefit)

                whats_new = item.get("whats_new") or []
                if whats_new:
                    whats = doc.add_paragraph()
                    wr = whats.add_run("What changed:")
                    wr.bold = True
                    if len(whats_new) == 1 and len(whats_new[0]) > 80:
                        doc.add_paragraph(whats_new[0])
                    else:
                        for bullet in whats_new:
                            doc.add_paragraph(bullet, style="List Bullet")

                doc.add_paragraph("")

        if payload["total_items"] == 0:
            doc.add_paragraph(
                "No tickets were moved to the Released section during the selected window."
            )

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
