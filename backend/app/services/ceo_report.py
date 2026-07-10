"""CEO Executive Brief email — one-minute leadership briefing, not a data dump."""

from __future__ import annotations

import io
from datetime import datetime, timedelta, timezone
from typing import Any, Literal

from docx import Document
from sqlalchemy.orm import Session

from app.config import get_settings
from app.services.activity_log import log_activity
from app.services.auto_sync import get_meta, _set_meta
from app.services.ceo_executive_briefing import apply_cursor_brief_overlay, build_executive_briefing
from app.services.ceo_intelligence import CEOIntelligenceService
from app.services.email_service import EmailService
from app.services.cursor_brief_store import load_ceo_brief_overlay

settings = get_settings()

ReportPeriod = Literal["weekly", "monthly", "6months"]

PERIOD_DAYS: dict[str, int] = {"weekly": 7, "monthly": 30, "6months": 180}

IST = timezone(timedelta(hours=5, minutes=30))

STATUS_STYLE = {
    "improving": ("Improving", "#047857", "#ecfdf5"),
    "stable": ("Stable", "#1d4ed8", "#eff6ff"),
    "watch": ("Watch", "#b45309", "#fffbeb"),
    "critical": ("Critical", "#b91c1c", "#fef2f2"),
}

AI_LABEL = {
    "positive": ("Positive", "#047857"),
    "mixed": ("Mixed", "#b45309"),
    "negative": ("Negative", "#b91c1c"),
    "insufficient": ("Insufficient Evidence", "#6b7280"),
}


def format_ist(iso_or_dt: str | datetime | None) -> str:
    if not iso_or_dt:
        return ""
    try:
        dt = datetime.fromisoformat(iso_or_dt.replace("Z", "+00:00")) if isinstance(iso_or_dt, str) else iso_or_dt
        if dt.tzinfo is None:
            dt = dt.replace(tzinfo=timezone.utc)
        return dt.astimezone(IST).strftime("%d %b %Y · %I:%M %p IST")
    except (ValueError, TypeError):
        return str(iso_or_dt)[:19]


def period_window(period: ReportPeriod, end: datetime | None = None) -> tuple[str, str, str]:
    end = end or datetime.utcnow()
    start = end - timedelta(days=PERIOD_DAYS[period])
    label = {
        "weekly": "Weekly Executive Brief",
        "monthly": "Monthly Executive Brief",
        "6months": "Six-Month Executive Brief",
    }[period]
    return start.strftime("%Y-%m-%d"), end.strftime("%Y-%m-%d"), label


def _esc(text: str) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _status_pill(status: str) -> str:
    label, color, bg = STATUS_STYLE.get(status, STATUS_STYLE["stable"])
    return (
        f'<span style="display:inline-block;padding:4px 12px;border-radius:6px;font-size:11px;'
        f'font-weight:600;color:{color};background:{bg}">{label}</span>'
    )


def _resolve_brief(db: Session, data: dict[str, Any], period: ReportPeriod) -> dict[str, Any]:
    brief = build_executive_briefing(data, period)
    overlay = load_ceo_brief_overlay(db)
    if overlay:
        brief = apply_cursor_brief_overlay(brief, overlay)
        brief["narrative_source"] = "cursor_weekly"
    else:
        brief["narrative_source"] = "rules"
    return brief


def build_report_html(data: dict[str, Any], period: ReportPeriod = "weekly", brief: dict[str, Any] | None = None) -> str:
    meta = data.get("meta") or {}
    if brief is None:
        brief = build_executive_briefing(data, period)
    generated = format_ist(datetime.utcnow())
    worry = brief.get("should_worry")

    # §1 CEO Brief — short headline
    brief_html = "".join(
        f'<p style="margin:0 0 10px;font-size:16px;line-height:1.55;color:#1c1917;font-weight:500">{_esc(s)}</p>'
        for s in brief.get("ceo_brief", [])
    )

    # §2 Scorecard — 3-column full-width grid with numbers + status reason
    scorecard_html = ""
    sc = brief.get("scorecard", [])
    for i in range(0, len(sc), 3):
        row_items = sc[i : i + 3]
        scorecard_html += "<tr>"
        for c in row_items:
            metric = c.get("metric") or c.get("sentence", "")
            reason = c.get("status_reason") or c.get("sentence", "")
            scorecard_html += f"""
        <td width="33%" style="padding:8px;vertical-align:top">
          <div style="background:#fafaf9;border:1px solid #e7e5e4;border-radius:10px;padding:18px 20px;height:100%">
            <table width="100%"><tr>
              <td style="font-size:14px;font-weight:600;color:#1c1917">{_esc(c['name'])}</td>
              <td style="text-align:right">{_status_pill(c.get('status', 'stable'))}</td>
            </tr></table>
            <p style="margin:12px 0 0;font-size:15px;font-weight:700;line-height:1.4;color:#0c0a09;font-variant-numeric:tabular-nums">{_esc(metric)}</p>
            <p style="margin:8px 0 0;font-size:12px;line-height:1.45;color:#78716c">{_esc(reason)}</p>
          </div>
        </td>"""
        for _ in range(3 - len(row_items)):
            scorecard_html += '<td width="33%"></td>'
        scorecard_html += "</tr>"

    # §3 Decisions — concise
    decisions_html = ""
    for d in brief.get("leadership_decisions", []):
        decisions_html += f"""
        <div style="margin:0 0 10px;padding:14px 18px;background:#1c1917;border-radius:8px">
          <p style="margin:0;font-size:14px;font-weight:600;color:#fafaf9;line-height:1.4">{_esc(d.get('decision', ''))}</p>
          <p style="margin:6px 0 0;font-size:12px;color:#a8a29e;line-height:1.4">{_esc(d.get('evidence', ''))}</p>
        </div>"""

    # §4 Evidence — chart + table
    ev = brief.get("evidence") or {}
    bars = ""
    for b in ev.get("chart") or []:
        bars += f"""
        <td align="center" style="vertical-align:bottom;padding:0 6px">
          <div style="width:32px;height:56px;display:flex;align-items:flex-end;background:#f5f5f4;border-radius:4px">
            <div style="width:100%;height:{b.get('pct', 10)}%;background:#f97316;border-radius:4px 4px 0 0"></div>
          </div>
          <p style="margin:6px 0 0;font-size:9px;color:#a8a29e">{_esc(b.get('month', ''))}</p>
        </td>"""

    table_rows = ""
    for row in ev.get("table") or []:
        dir_color = "#b91c1c" if row.get("direction") == "Rising" else "#047857" if row.get("direction") == "Falling" else "#57534e"
        table_rows += f"""
        <tr>
          <td style="padding:10px 12px;border-bottom:1px solid #f5f5f4;font-size:12px;color:#1c1917">{_esc(row.get('area', ''))}</td>
          <td style="padding:10px 12px;border-bottom:1px solid #f5f5f4;font-size:12px;text-align:center;color:#78716c">{_esc(row.get('before', ''))}</td>
          <td style="padding:10px 12px;border-bottom:1px solid #f5f5f4;font-size:12px;text-align:center;font-weight:600;color:#1c1917">{_esc(row.get('current', ''))}</td>
          <td style="padding:10px 12px;border-bottom:1px solid #f5f5f4;font-size:12px;text-align:center;color:{dir_color};font-weight:600">{_esc(row.get('direction', ''))}</td>
        </tr>"""

    ai = brief.get("ai_impact") or {}
    verdict = ai.get("verdict", "insufficient")
    v_label, v_color = AI_LABEL.get(verdict, AI_LABEL["insufficient"])

    questions_html = "".join(
        f'<li style="margin:0 0 8px;font-size:13px;line-height:1.45;color:#1c1917">{_esc(q)}</li>'
        for q in brief.get("leadership_questions", [])
    )

    mood = (
        '<span style="color:#b91c1c">Attention needed</span>'
        if worry
        else '<span style="color:#047857">Stable</span>'
    )

    return f"""<!DOCTYPE html>
<html lang="en"><head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<meta name="color-scheme" content="light dark">
<style>@media (prefers-color-scheme:dark){{.card{{background:#171717!important}}.body-text{{color:#e7e5e4!important}}}}</style>
</head>
<body style="margin:0;padding:0;background:#f5f5f4;font-family:Georgia,'Times New Roman',-apple-system,BlinkMacSystemFont,sans-serif">
<table width="100%" cellpadding="0" cellspacing="0" style="background:#f5f5f4"><tr><td>

<table class="card" width="100%" cellpadding="0" cellspacing="0" style="width:100%;background:#ffffff;border-collapse:collapse">

<tr><td style="background:#0c0a09;padding:32px 40px 28px">
  <p style="margin:0;font-size:11px;letter-spacing:.18em;text-transform:uppercase;color:#a8a29e">Autorox · Executive Intelligence</p>
  <h1 style="margin:12px 0 0;font-size:26px;font-weight:400;color:#fafaf9;letter-spacing:-.01em">{_esc(brief.get('period_label', 'Executive Brief'))}</h1>
  <p style="margin:10px 0 0;font-size:12px;color:#78716c">{generated} · Overall: {mood} · Health {data.get('health_score', {}).get('score', '—')}/100</p>
</td></tr>

<tr><td style="padding:32px 40px 40px">

<p style="margin:0 0 14px;font-size:10px;font-weight:700;letter-spacing:.14em;text-transform:uppercase;color:#a8a29e">CEO Brief</p>
<div style="padding:22px 24px;background:#fafaf9;border-radius:10px;border:1px solid #e7e5e4;margin-bottom:32px">
{brief_html}
</div>

<p style="margin:0 0 14px;font-size:10px;font-weight:700;letter-spacing:.14em;text-transform:uppercase;color:#a8a29e">Executive Scorecard</p>
<table width="100%" cellpadding="0" cellspacing="0" style="margin-bottom:32px">{scorecard_html}</table>

<p style="margin:0 0 14px;font-size:10px;font-weight:700;letter-spacing:.14em;text-transform:uppercase;color:#a8a29e">Leadership Decisions</p>
{decisions_html or '<p style="font-size:12px;color:#a8a29e">No decisions required this period.</p>'}

<p style="margin:32px 0 14px;font-size:10px;font-weight:700;letter-spacing:.14em;text-transform:uppercase;color:#a8a29e">Supporting Evidence</p>
<p style="margin:0 0 10px;font-size:12px;color:#57534e;font-style:italic">{_esc(ev.get('chart_title', ''))}</p>
<table width="100%" cellpadding="0" cellspacing="0" style="margin-bottom:24px"><tr>{bars}</tr></table>
<p style="margin:0 0 8px;font-size:12px;color:#57534e">{_esc(ev.get('table_title', ''))}</p>
<table width="100%" cellpadding="0" cellspacing="0" style="border:1px solid #e7e5e4;border-radius:8px;overflow:hidden;margin-bottom:28px">
  <tr style="background:#fafaf9">
    <th style="padding:12px 16px;text-align:left;font-size:10px;text-transform:uppercase;color:#a8a29e;font-weight:600">Area</th>
    <th style="padding:12px;text-align:center;font-size:10px;text-transform:uppercase;color:#a8a29e">Before adoption</th>
    <th style="padding:12px;text-align:center;font-size:10px;text-transform:uppercase;color:#a8a29e">Current</th>
    <th style="padding:12px;text-align:center;font-size:10px;text-transform:uppercase;color:#a8a29e">Trend</th>
  </tr>
  {table_rows}
</table>

<p style="margin:0 0 14px;font-size:10px;font-weight:700;letter-spacing:.14em;text-transform:uppercase;color:#a8a29e">Questions for Leadership</p>
<ol style="margin:0;padding-left:20px">{questions_html}</ol>

<p style="margin:32px 0 0;padding-top:16px;border-top:1px solid #e7e5e4;font-size:10px;color:#a8a29e;line-height:1.6">
  Prepared for CEO review · Autorox Delivery Intelligence · Correlational analysis only
  {f" · Narrative drafted by Cursor AI from verified metrics" if brief.get("narrative_source") == "cursor_weekly" else ""}
</p>

</td></tr>
</table>
</td></tr></table>
</body></html>"""


def build_report_text(data: dict[str, Any], period: ReportPeriod = "weekly", brief: dict[str, Any] | None = None) -> str:
    if brief is None:
        brief = build_executive_briefing(data, period)
    lines = [brief.get("subject", "Executive Brief"), "=" * 50, "", "CEO BRIEF", ""]
    lines.extend(brief.get("ceo_brief", []))
    lines.extend(["", "EXECUTIVE SCORECARD"])
    for c in brief.get("scorecard", []):
        metric = c.get("metric") or c.get("sentence", "")
        reason = c.get("status_reason") or ""
        lines.append(f"  {c['name']}: {c.get('status', '').upper()}")
        lines.append(f"    {metric}")
        if reason:
            lines.append(f"    → {reason}")
    lines.extend(["", "LEADERSHIP DECISIONS"])
    for d in brief.get("leadership_decisions", []):
        lines.append(f"  → {d.get('decision')}")
        lines.append(f"    {d.get('evidence')}")
    lines.extend(["", "QUESTIONS FOR LEADERSHIP"])
    for q in brief.get("leadership_questions", []):
        lines.append(f"  ? {q}")
    return "\n".join(lines)


def build_report_docx(data: dict[str, Any], period: ReportPeriod = "weekly", brief: dict[str, Any] | None = None) -> bytes:
    if brief is None:
        brief = build_executive_briefing(data, period)
    doc = Document()
    doc.add_paragraph(brief.get("period_label", "Executive Brief")).bold = True
    doc.add_paragraph("")
    doc.add_paragraph("CEO Brief").bold = True
    for s in brief.get("ceo_brief", []):
        doc.add_paragraph(s)
    doc.add_paragraph("")
    doc.add_paragraph("Leadership Decisions").bold = True
    for d in brief.get("leadership_decisions", []):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(d.get("decision", "")).bold = True
        doc.add_paragraph(d.get("evidence", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class CEOReportService:
    META_LAST_SENT = "ceo_report_last_sent_at"
    META_SCHEDULE_ENABLED = "ceo_report_schedule_enabled"
    META_SCHEDULE_FREQUENCY = "ceo_report_schedule_frequency"
    META_SCHEDULE_PROJECT = "ceo_report_schedule_project_gid"
    META_CEO_EMAIL = "ceo_report_ceo_email"

    def __init__(self, db: Session):
        self.db = db
        self.email = EmailService()

    def get_settings(self) -> dict[str, Any]:
        from app.services.weekly_analysis_pipeline import get_weekly_analysis_status

        base = {
            "ceo_email": get_meta(self.db, self.META_CEO_EMAIL) or settings.ceo_email,
            "ai_adoption_date": settings.ai_adoption_date,
            "schedule_enabled": (get_meta(self.db, self.META_SCHEDULE_ENABLED) or "false").lower() == "true",
            "schedule_frequency": get_meta(self.db, self.META_SCHEDULE_FREQUENCY) or settings.ceo_report_frequency,
            "schedule_project_gid": get_meta(self.db, self.META_SCHEDULE_PROJECT),
            "last_sent_at": get_meta(self.db, self.META_LAST_SENT),
            "email_configured": self.email.configured,
            "cursor_configured": settings.cursor_configured,
            "weekly_analysis": get_weekly_analysis_status(self.db),
            "schedule_note": (
                "Issue Intelligence + Cursor enrich runs daily at 5 PM IST. "
                "When enabled: CEO email sends Tuesday ~8 AM IST."
            ),
        }
        return base

    def update_settings(
        self,
        *,
        ceo_email: str | None = None,
        schedule_enabled: bool | None = None,
        schedule_frequency: str | None = None,
        schedule_project_gid: str | None = None,
    ) -> dict[str, Any]:
        if ceo_email is not None:
            _set_meta(self.db, self.META_CEO_EMAIL, ceo_email.strip().lower())
        if schedule_enabled is not None:
            _set_meta(self.db, self.META_SCHEDULE_ENABLED, "true" if schedule_enabled else "false")
        if schedule_frequency is not None:
            _set_meta(self.db, self.META_SCHEDULE_FREQUENCY, schedule_frequency)
        if schedule_project_gid is not None:
            _set_meta(self.db, self.META_SCHEDULE_PROJECT, schedule_project_gid or "")
        return self.get_settings()

    def build_intelligence(
        self,
        project_gid: str | None,
        period: ReportPeriod,
    ) -> tuple[dict[str, Any], str, str, str]:
        date_from, date_to, label = period_window(period)
        data = CEOIntelligenceService(
            self.db, project_gid, date_from=date_from, date_to=date_to
        ).build()
        return data, date_from, date_to, label

    def preview_report(
        self,
        *,
        project_gid: str | None,
        period: ReportPeriod = "weekly",
    ) -> dict[str, Any]:
        data, date_from, date_to, label = self.build_intelligence(project_gid, period)
        brief = _resolve_brief(self.db, data, period)
        return {
            "subject": brief.get("subject", "CEO Weekly Engineering Brief"),
            "period": period,
            "period_label": brief.get("period_label", label),
            "date_from": date_from,
            "date_to": date_to,
            "health_score": int((data.get("health_score") or {}).get("score", 0)),
            "narrative_source": brief.get("narrative_source", "rules"),
            "cursor_generated_at": brief.get("cursor_generated_at"),
            "html": build_report_html(data, period, brief=brief),
            "text": build_report_text(data, period, brief=brief),
            "brief": brief,
        }

    def send_report(
        self,
        *,
        project_gid: str | None,
        period: ReportPeriod = "weekly",
        recipient_emails: list[str] | None = None,
        source: str = "manual",
    ) -> dict[str, Any]:
        if not self.email.configured:
            raise RuntimeError("Gmail not configured. Set GMAIL_USER and GMAIL_APP_PASSWORD in .env")

        emails = [e.strip().lower() for e in (recipient_emails or []) if e and e.strip()]
        if not emails:
            emails = [get_meta(self.db, self.META_CEO_EMAIL) or settings.ceo_email]
        emails = sorted(set(emails))

        data, date_from, date_to, _label = self.build_intelligence(project_gid, period)
        brief = _resolve_brief(self.db, data, period)
        subject = brief.get("subject", "CEO Weekly Engineering Brief")

        text = build_report_text(data, period, brief=brief)
        html = build_report_html(data, period, brief=brief)
        docx = build_report_docx(data, period, brief=brief)
        filename = f"Executive Brief {date_to}.docx"

        self.email.send_email(emails, subject, text, html, attachment=(filename, docx))

        now = datetime.utcnow().isoformat()
        _set_meta(self.db, self.META_LAST_SENT, now)

        log_activity(
            self.db,
            module="ceo_intelligence",
            action="report_sent",
            summary=f"Sent {period} executive brief to {', '.join(emails)}",
            entity_type="ceo_report",
            entity_id=period,
            payload={
                "recipients": emails,
                "period": period,
                "date_from": date_from,
                "date_to": date_to,
                "subject": subject,
                "source": source,
                "narrative_source": brief.get("narrative_source"),
            },
        )

        return {
            "success": True,
            "recipient_count": len(emails),
            "sent_to": emails,
            "period": period,
            "date_from": date_from,
            "date_to": date_to,
            "health_score": data.get("health_score", {}).get("score"),
            "sent_at": now,
        }
